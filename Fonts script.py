from matplotlib import font_manager
from docx import Document
from docx.shared import Pt
import unicodedata

# Create a new Word document
doc = Document()
# Add a table to the document. The table will have 2 columns and as many rows as there are fonts
table = doc.add_table(rows=1, cols=2)
# Set the header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Font Name'
hdr_cells[1].text = 'Font Sample'

# Get all installed fonts
fonts = font_manager.findSystemFonts(fontpaths=None, fontext='ttf')

# Define the sample text to display for each font, including Hebrew characters and numbers
sample_text = 'אבגדהוזחטיכלמנסעפצקרשת 0123456789'
hebrew_char_range = (0x0590, 0x05FF)  # Unicode range for Hebrew characters

# Filter out non-Hebrew supporting fonts
fonts_with_hebrew = []
for font in fonts:
    font_prop = font_manager.FontProperties(fname=font)
    if any(unicodedata.category(chr(cp)).startswith('Lo') and
           hebrew_char_range[0] <= cp <= hebrew_char_range[1]
           for cp in range(hebrew_char_range[0], hebrew_char_range[1]+1)):
        fonts_with_hebrew.append(font_prop.get_name())

# Populate the table with font names and samples
for font_name in fonts_with_hebrew:
    row_cells = table.add_row().cells
    row_cells[0].text = font_name
    run = row_cells[1].paragraphs[0].add_run(sample_text)
    run.font.name = font_name
    run.font.size = Pt(12)

# Save the document
doc.save('font_samples.docx')

print('Word document with font samples has been created.')
